#!/usr/bin/env python3
"""Ingest source material into markdown.

Supports: .md/.txt (copy), .docx (python-docx), .pdf (markitdown cascade).
URLs and pastes are handled by the orchestrating agent, not this script.

Usage:
    ingest.py <source> <out.md>
    ingest.py <source_dir>/  <out_dir>/     # batch: one .md per source file

PDF cascade: uvx markitdown -> pipx run markitdown -> pypdf (uvx) -> fail loud.
"""
import os
import re
import sys
import shutil
import subprocess


def _docx_to_md(path):
    import docx
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        style = (p.style.name or '').lower()
        txt = p.text.rstrip()
        if not txt:
            parts.append('')
            continue
        if 'heading 1' in style or style == 'title':
            parts.append(f'# {txt}')
        elif 'heading 2' in style:
            parts.append(f'## {txt}')
        elif 'heading 3' in style:
            parts.append(f'### {txt}')
        elif 'heading 4' in style:
            parts.append(f'#### {txt}')
        elif 'list' in style:
            parts.append(f'- {txt}')
        else:
            parts.append(txt)
    for t in d.tables:
        parts.append('')
        rows = []
        for row in t.rows:
            cells = [c.text.strip().replace('|', '\\|') for c in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
        if rows:
            sep = '|' + '|'.join(['---'] * len(t.columns)) + '|'
            parts.append(rows[0])
            parts.append(sep)
            parts.extend(rows[1:])
    return '\n'.join(parts).strip() + '\n'


def _pdf_cascade(path):
    for cmd in (['uvx', 'markitdown', path], ['pipx', 'run', 'markitdown', path]):
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
            if r.returncode == 0 and r.stdout.strip():
                return r.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    # last resort: pypdf via uvx one-shot
    code = (
        "import sys,pypdf; "
        "r=pypdf.PdfReader(sys.argv[1]); "
        "print('\\n\\n'.join((p.extract_text() or '') for p in r.pages))"
    )
    try:
        r = subprocess.run(['uvx', '--from', 'pypdf', 'python', '-c', code, path],
                           capture_output=True, text=True, timeout=120)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    raise RuntimeError(
        f"PDF ingestion failed for {path}. Install markitdown (`pipx install markitdown`) "
        "or pypdf, or paste the text manually.")


def ingest_one(src, out_path):
    ext = os.path.splitext(src)[1].lower()
    if ext in ('.md', '.markdown', '.txt'):
        shutil.copyfile(src, out_path)
    elif ext == '.docx':
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(_docx_to_md(src))
    elif ext == '.pdf':
        md = _pdf_cascade(src)
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(md)
    else:
        raise RuntimeError(f"unsupported source type: {ext}")
    return out_path


def ingest_batch(src_dir, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    written = []
    for name in sorted(os.listdir(src_dir)):
        src = os.path.join(src_dir, name)
        if not os.path.isfile(src):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext not in ('.md', '.markdown', '.txt', '.docx', '.pdf'):
            continue
        base = os.path.splitext(name)[0]
        out = os.path.join(out_dir, re.sub(r'\W+', '_', base) + '.md')
        ingest_one(src, out)
        written.append(out)
    return written


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("usage: ingest.py <source_file_or_dir> <out_file_or_dir>", file=sys.stderr)
        sys.exit(2)
    src, out = sys.argv[1], sys.argv[2]
    if os.path.isdir(src):
        for w in ingest_batch(src, out):
            print("wrote", w)
    else:
        ingest_one(src, out)
        print("wrote", out)
